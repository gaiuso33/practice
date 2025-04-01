import pandas as pd
from docx import Document
from docx.shared import Inches

# 1️⃣ Load the Excel file
file_path = "your_file.xls"  # Change to your actual file path
df = pd.read_excel(file_path, sheet_name="Sheet1")  # Adjust sheet name if necessary

# 2️⃣ Select only 'Project Title' and 'Director Name' columns
df_selected = df[['Project Title', 'Director Name']]

# 3️⃣ Group multiple projects under the same director
df_grouped = df_selected.groupby('Director Name')['Project Title'].apply(', '.join).reset_index()

# 4️⃣ Create a Word document
doc = Document()
doc.add_heading('Project List by Director', level=1)  # Title

# 5️⃣ Add a table
table = doc.add_table(rows=1, cols=2)  # Create table with 2 columns
table.style = 'Table Grid'  # Add borders

# 6️⃣ Add column headers
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Director Name'
hdr_cells[1].text = 'Projects'

# 7️⃣ Populate the table with data
for index, row in df_grouped.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = row['Director Name']
    row_cells[1].text = row['Project Title']

# 8️⃣ Save the document
word_file = "grouped_projects.docx"
doc.save(word_file)

print(f"✅ Word document '{word_file}' with a table has been created successfully!")
