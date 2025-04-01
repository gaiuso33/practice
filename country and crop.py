import pandas as pd
from docx import Document

# 1️⃣ Load the Excel file
file_path = "C:\\Users\\Ayogaius\\Downloads\\2025-03-27 List of active projects.xls"  # Change to your actual file path
df = pd.read_excel(file_path, sheet_name="List of active projects", engine="xlrd")
  # Adjust sheet name if necessary

# 2️⃣ Select relevant columns (Assuming 'Country' and 'Crop' exist)
df_selected = df[['Project', 'Country', 'Crop']] 

### 🔹 Function to Save Grouped Data to a Word Table (List Format) ###
def save_to_word(df_grouped, file_name, category_name):
    doc = Document()
    doc.add_heading(f'Projects Sorted by {category_name}', level=1)

    # Create a table with 2 columns: Category & Projects
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Add column headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = category_name
    hdr_cells[1].text = 'Projects'

    # Populate the table
    for index, row in df_grouped.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row[category_name])
        
        # Convert project list into a bullet point format
        project_list = row['Project'].split('\n')
        row_cells[1].text = "\n".join(f"- {proj}" for proj in project_list)

    # Save the document
    doc.save(file_name)
    print(f"✅ Word document '{file_name}' has been created successfully!")

# 3️⃣ Group by 'Country' and list projects
df_grouped_country = df_selected.groupby('Country')['Project'].apply(lambda x: '\n'.join(x)).reset_index()
save_to_word(df_grouped_country, "projects_by_country.docx", "Country")

# 4️⃣ Group by 'Crop' and list projects
df_grouped_crop = df_selected.groupby('Crop')['Project'].apply(lambda x: '\n'.join(x)).reset_index()
save_to_word(df_grouped_crop, "projects_by_crop.docx", "Crop")